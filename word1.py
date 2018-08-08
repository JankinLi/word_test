import docx;

def test1():
    print("test1 begin.");
    #print("""E:\lichuan\势坤科技\examples""");
    filePath = """E:\lichuan\势坤科技\examples\成都电子科技大学信息表.docx""";
    print(filePath);
    file = docx.Document(filePath);
    print("段落数:" + str(len(file.paragraphs)));
    tables = file.tables;
    print("表格数："+ str(len(tables)));
    for table in tables:
        printTable(table);

def printTable(table):
    # 输出每一表格的内容
    lineNo=0;
    for row in table.rows:
        lineNo=lineNo+1;
        cells = row.cells;
        colNo = 0;
        savedText = "";
        for cell in cells:
            if( savedText == cell.text):
                continue;
            colNo = colNo + 1;
            print("line:"+ str(lineNo) + ",col:" + str(colNo) + "," + cell.text);
            savedText = cell.text;


    # 输出段落编号及段落内容
    #for i in range(len(file.paragraphs)):
    #    print("第" + str(i) + "段的内容是：" + file.paragraphs[i].text)
    print("test1 end.");


if __name__ == "__main__":
    print("word test.");
    test1();

